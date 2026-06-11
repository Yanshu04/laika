import time
import docx
import httpx
from pathlib import Path

BASE_URL = "http://127.0.0.1:8000"

def create_dummy_docx(file_path: Path):
    """Create a temporary docx file with specific mock facts."""
    doc = docx.Document()
    doc.add_heading("LAIKA Test Manual", 0)
    doc.add_paragraph(
        "This is the official test manual for the LAIKA Knowledge Assistant. "
        "The official company color is neon-glass. "
        "The company CEO is Antigravity AI."
    )
    doc.save(str(file_path))
    print(f"Created temporary test file at: {file_path}")

def run_test():
    # Setup paths
    temp_file = Path(__file__).parent / "laika_test_manual.docx"
    create_dummy_docx(temp_file)
    
    try:
        # 1. Upload the file
        print("Uploading test document...")
        with open(temp_file, "rb") as f:
            files = {"files": ("laika_test_manual.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
            response = httpx.post(f"{BASE_URL}/api/documents/upload", files=files, timeout=15.0)
            
        assert response.status_code == 202, f"Upload failed: {response.text}"
        doc_data = response.json()[0]
        doc_id = doc_data["id"]
        print(f"Document uploaded. ID: {doc_id}. Status: {doc_data['status']}")
        
        # 2. Wait for indexing to complete
        print("Waiting for indexing to complete...")
        indexed = False
        for _ in range(30):  # Wait up to 30 seconds
            time.sleep(1.0)
            status_res = httpx.get(f"{BASE_URL}/api/documents", timeout=5.0)
            assert status_res.status_code == 200
            docs = status_res.json()
            doc_status = next((d["status"] for d in docs if d["id"] == doc_id), None)
            
            if doc_status == "INDEXED":
                print("Document successfully indexed!")
                indexed = True
                break
            elif doc_status == "ERROR":
                print("Document indexing failed with status ERROR.")
                break
            else:
                print(f"Current status: {doc_status}...")
                
        assert indexed, "Indexing timed out or failed."
        
        # 3. Query the chat pipeline
        print("Querying the RAG chat engine...")
        query_payload = {"question": "What is the company color and who is the CEO?"}
        query_res = httpx.post(f"{BASE_URL}/api/chat/query", json=query_payload, timeout=60.0)
        
        assert query_res.status_code == 200, f"Query failed: {query_res.text}"
        data = query_res.json()
        
        print("\n--- LLM ANSWER ---")
        print(data["answer"])
        print("------------------")
        
        print("\n--- SOURCE CITATIONS ---")
        for src in data["sources"]:
            print(f"- File: {src['filename']}, Snippet: \"{src['text_snippet']}\"")
        print("------------------------\n")
        
        # Verify LLM has answer correctly
        assert "neon-glass" in data["answer"].lower() or "neon-glass" in data["answer"], "Failed to find correct company color in answer"
        assert "antigravity" in data["answer"].lower(), "Failed to find correct CEO in answer"
        assert len(data["sources"]) >= 1, "Expected at least one source citation"
        print("RAG query test passed successfully!")
        
        # 4. Delete the document
        print("Cleaning up: Deleting document...")
        delete_res = httpx.delete(f"{BASE_URL}/api/documents/{doc_id}", timeout=5.0)
        assert delete_res.status_code == 200, f"Deletion failed: {delete_res.text}"
        print("Cleanup completed successfully.")
        
    finally:
        # Delete temp file
        if temp_file.exists():
            temp_file.unlink()
            print("Removed temporary test file.")

if __name__ == "__main__":
    run_test()
