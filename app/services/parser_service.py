import docx
from pathlib import Path
from typing import List, Dict, Any, Optional
from pypdf import PdfReader
from app.utils.logger import get_logger

logger = get_logger(__name__)

class PDFParser:
    @staticmethod
    def parse(file_path: Path) -> List[Dict[str, Any]]:
        """
        Parses a PDF file page by page.
        Returns a list of dicts: [{"page_number": int, "text": str}]
        """
        logger.info(f"Parsing PDF file: {file_path.name}")
        pages = []
        try:
            reader = PdfReader(file_path)
            total_pages = len(reader.pages)
            logger.info(f"PDF has {total_pages} pages")
            
            for i, page in enumerate(reader.pages):
                text = page.extract_text() or ""
                pages.append({
                    "page_number": i + 1,
                    "text": text.strip()
                })
        except Exception as e:
            logger.error(f"Error parsing PDF {file_path.name}: {e}")
            raise RuntimeError(f"PDF parsing failed: {e}") from e
            
        return pages

class DOCXParser:
    @staticmethod
    def parse(file_path: Path) -> List[Dict[str, Any]]:
        """
        Parses a DOCX file, extracting paragraphs and tables.
        Returns a list of dicts: [{"page_number": None, "text": str}]
        """
        logger.info(f"Parsing DOCX file: {file_path.name}")
        text_elements = []
        try:
            doc = docx.Document(str(file_path))
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                cleaned_text = paragraph.text.strip()
                if cleaned_text:
                    text_elements.append(cleaned_text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_cells_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    # De-duplicate cell text that might span multiple cells
                    seen = set()
                    unique_cells = []
                    for val in row_cells_text:
                        if val not in seen:
                            seen.add(val)
                            unique_cells.append(val)
                    if unique_cells:
                        text_elements.append(" | ".join(unique_cells))
            
            combined_text = "\n".join(text_elements)
            
            return [{
                "page_number": None,
                "text": combined_text.strip()
            }]
        except Exception as e:
            logger.error(f"Error parsing DOCX {file_path.name}: {e}")
            raise RuntimeError(f"DOCX parsing failed: {e}") from e

class TextParser:
    @staticmethod
    def parse(file_path: Path) -> List[Dict[str, Any]]:
        """
        Parses a plain text or Markdown file.
        Returns a list of dicts: [{"page_number": None, "text": str}]
        """
        logger.info(f"Parsing Text/Markdown file: {file_path.name}")
        try:
            try:
                with open(file_path, "r", encoding="utf-8") as f:
                    content = f.read()
            except UnicodeDecodeError:
                with open(file_path, "r", encoding="latin-1") as f:
                    content = f.read()
            return [{
                "page_number": None,
                "text": content.strip()
            }]
        except Exception as e:
            logger.error(f"Error parsing Text/Markdown file {file_path.name}: {e}")
            raise RuntimeError(f"Text parsing failed: {e}") from e

class TextExtractionService:
    @staticmethod
    def extract_text(file_path: Path) -> List[Dict[str, Any]]:
        """
        Detect file type and extract pages/sections with text.
        """
        suffix = file_path.suffix.lower()
        if suffix == ".pdf":
            return PDFParser.parse(file_path)
        elif suffix == ".docx":
            return DOCXParser.parse(file_path)
        elif suffix in [".txt", ".md"]:
            return TextParser.parse(file_path)
        else:
            logger.error(f"Unsupported file format: {suffix}")
            raise ValueError(f"Unsupported file extension: {suffix}")
